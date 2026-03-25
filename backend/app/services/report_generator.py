"""보고서 생성 서비스 — AI 분석 결과를 PDF/DOCX 풀 보고서로 변환"""

from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from app.models.ai_analysis import AIAnalysis
    from app.models.startup import Startup

logger = logging.getLogger(__name__)

TYPE_LABEL = {
    "screening": "AI 스크리닝",
    "ir_analysis": "IR 심층분석",
    "risk_alert": "리스크 스캔",
    "market_scan": "시장 분석",
}

SCORE_LABEL = {
    "team_competency": "팀 역량",
    "tech_differentiation": "제품/기술",
    "market_potential": "시장 규모 및 성장성",
    "problem_clarity": "문제 정의 명확성",
    "initial_validation": "초기 검증",
    "strategy_fit": "전략 적합성",
    "business_model": "비즈니스 모델",
    "financial_health": "재무 건전성",
    "growth_potential": "성장성",
    "team_strength": "팀 구성",
    "ip_competitiveness": "IP 경쟁력",
    "team_risk": "팀 리스크",
    "financial_risk": "재무 리스크",
    "tech_risk": "기술 리스크",
    "market_risk": "시장 리스크",
    "legal_risk": "법적 리스크",
    "market_size": "시장 규모",
    "growth_rate": "성장률",
    "competition_intensity": "경쟁 강도",
    "entry_barrier": "진입 장벽",
    "regulatory_environment": "규제 환경",
    "vision_scalability": "비전 및 확장성",
    "solution_differentiation": "솔루션 차별화",
    "investment_fit": "투자 적합성",
    "presentation": "프리젠테이션",
}

REC_LABEL = {"pass": "통과", "conditional": "조건부 통과", "hold": "보류", "decline": "거절"}
RISK_LABEL = {"low": "낮음", "medium": "보통", "high": "높음", "critical": "심각"}


def _safe_json(summary: str) -> dict[str, Any]:
    """summary 필드에 포함된 JSON 구조 데이터를 추출. 없으면 빈 dict."""
    # scores 필드에 full_report가 저장되어 있을 수 있음
    return {}


def _parse_sections(summary: str) -> tuple[str, list[str], str]:
    """summary를 (AI 분석 본문, 핵심 발견 리스트, 규칙 기반 분석)으로 분리"""
    import re
    sections = summary.split("---")
    llm_part = sections[0].strip() if sections else summary
    rule_part = sections[1].strip() if len(sections) > 1 else ""

    findings: list[str] = []
    main_text = llm_part
    if "핵심 발견:" in llm_part:
        idx = llm_part.index("핵심 발견:")
        main_text = llm_part[:idx].strip()
        findings_block = llm_part[idx + len("핵심 발견:"):].strip()
        for line in findings_block.split("\n"):
            cleaned = line.strip().lstrip("•-").strip()
            if cleaned:
                findings.append(cleaned)

    for prefix in ["📄 문서 기반 AI 분석", "📊 규칙 기반 분석"]:
        main_text = main_text.replace(prefix, "").strip()
        rule_part = rule_part.replace(prefix, "").strip()

    main_text = re.sub(r"\(\d+건 분석\)\s*", "", main_text).strip()
    rule_part = re.sub(r"\(참고\)\s*", "", rule_part).strip()

    return main_text, findings, rule_part


def _extract_full_report(analysis: AIAnalysis) -> dict[str, Any] | None:
    """scores JSON에서 풀 보고서 구조 데이터를 추출"""
    scores = analysis.scores
    if not scores:
        return None
    # 풀 보고서 구조가 scores에 저장된 경우 (score_details 키 존재)
    if "score_details" in scores or "top_strengths" in scores:
        return scores
    return None


# ── PDF 생성 ──────────────────────────────────────────────────────

def generate_pdf(analysis: AIAnalysis, startup: Startup) -> bytes:
    """AI 분석 결과를 풀 보고서 PDF로 생성"""
    from fpdf import FPDF
    import glob
    import os

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # 나눔고딕 폰트
    font_paths = glob.glob("/usr/share/fonts/truetype/nanum/NanumGothic.ttf")
    bold_path = ""
    if font_paths:
        pdf.add_font("NG", "", font_paths[0], uni=True)
        bold_candidate = font_paths[0].replace("NanumGothic.ttf", "NanumGothicBold.ttf")
        if os.path.exists(bold_candidate):
            pdf.add_font("NG", "B", bold_candidate, uni=True)
            bold_path = bold_candidate
        else:
            pdf.add_font("NG", "B", font_paths[0], uni=True)
        fn = "NG"
    else:
        fn = "Helvetica"

    type_label = TYPE_LABEL.get(analysis.analysis_type, analysis.analysis_type)
    rec_label = REC_LABEL.get(analysis.recommendation or "", analysis.recommendation or "-")
    risk_label = RISK_LABEL.get(analysis.risk_level or "", analysis.risk_level or "-")
    created = analysis.created_at.strftime("%Y.%m.%d") if analysis.created_at else datetime.now().strftime("%Y.%m.%d")

    scores = analysis.scores or {}
    main_text, findings, rule_part = _parse_sections(analysis.summary or "")

    # 풀 보고서 데이터 추출 시도
    full = _extract_full_report(analysis)
    score_details = full.get("score_details", []) if full else []
    top_strengths = full.get("top_strengths", []) if full else []
    top_concerns = full.get("top_concerns", []) if full else []
    verification_needed = full.get("verification_needed", []) if full else []
    overall_opinion = full.get("overall_opinion", "") if full else ""
    company_overview = full.get("company_overview", {}) if full else {}
    eligibility = full.get("eligibility_check", {}) if full else {}
    one_line = full.get("one_line_summary", "") if full else ""
    web_research = scores.get("_web_research", "")

    # 점수 계산
    score_items = {k: v for k, v in scores.items() if isinstance(v, (int, float))}
    avg_score = sum(score_items.values()) / len(score_items) if score_items else 0
    sorted_scores = sorted(score_items.items(), key=lambda x: x[1], reverse=True)

    def heading(text: str, size: int = 13):
        pdf.ln(3)
        pdf.set_font(fn, "B", size)
        pdf.cell(0, 8, text, ln=True)
        pdf.ln(2)

    def body(text: str, size: int = 10):
        pdf.set_font(fn, "", size)
        pdf.multi_cell(0, 6, text)

    def separator():
        pdf.ln(3)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)

    # ── 표지 ──
    pdf.set_font(fn, "B", 20)
    pdf.cell(0, 14, f"{type_label} 보고서", ln=True, align="C")
    pdf.ln(2)
    pdf.set_font(fn, "", 11)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 7, f"{startup.company_name}", ln=True, align="C")
    pdf.cell(0, 7, f"작성일: {created}  |  작성자: AI 에이전트", ln=True, align="C")
    pdf.cell(0, 7, f"{startup.industry or '-'}  |  {startup.stage or '-'}", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)
    pdf.set_font(fn, "", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, "본 보고서는 AI 에이전트의 분석을 기반으로 작성되었으며, 담당자의 검토를 거쳐 확정되어야 합니다.", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    separator()

    # ── 1. 요약 ──
    heading("1. 요약")
    pdf.set_font(fn, "", 10)
    col_w = 47.5
    pdf.cell(col_w, 7, f"종합 점수: {avg_score:.1f}/5", border=1, align="C")
    pdf.cell(col_w, 7, f"판정: {rec_label}", border=1, align="C")
    pdf.cell(col_w, 7, f"리스크: {risk_label}", border=1, align="C")
    pdf.cell(col_w, 7, f"매력도: {analysis.investment_attractiveness or '-'}/5", border=1, align="C")
    pdf.ln(10)

    if one_line:
        body(f"한줄 요약: {one_line}")
        pdf.ln(3)
    elif main_text:
        body(main_text[:500])
        pdf.ln(3)

    separator()

    # ── 2. 기본 적격성 검사 ──
    if eligibility:
        heading("2. 기본 적격성 검사")
        pdf.set_font(fn, "B", 9)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(70, 7, "항목", border=1, fill=True)
        pdf.cell(30, 7, "결과", border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font(fn, "", 9)
        labels = {"legal_issue": "법적 결격사유", "duplicate_application": "중복 지원", "antisocial_model": "반사회적 사업", "portfolio_conflict": "포트폴리오 충돌"}
        for key, label in labels.items():
            val = eligibility.get(key, "미확인")
            pdf.cell(70, 7, f"  {label}", border=1)
            pdf.cell(30, 7, val, border=1, align="C")
            pdf.ln()
        separator()

    # ── 3. 기업 개요 ──
    if company_overview:
        sec_num = 3
        heading(f"{sec_num}. 기업 개요")
        pdf.set_font(fn, "", 9)
        for key, val in company_overview.items():
            if val:
                label = {"company_name": "기업명", "representative": "대표자", "business_field": "사업 분야", "main_product": "주요 제품", "team_size": "팀 규모", "investment_request": "투자 요청", "government_rd": "정부 R&D"}.get(key, key)
                pdf.cell(40, 6, f"  {label}", border=0)
                pdf.cell(0, 6, str(val), border=0)
                pdf.ln()
        separator()

    # ── 4. 항목별 평가 ──
    heading("4. 항목별 점수")

    # 점수 테이블
    pdf.set_font(fn, "B", 9)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(60, 7, "평가 항목", border=1, fill=True)
    pdf.cell(25, 7, "점수", border=1, fill=True, align="C")
    pdf.cell(25, 7, "등급", border=1, fill=True, align="C")
    pdf.ln()
    pdf.set_font(fn, "", 9)
    for key, value in sorted_scores:
        label = SCORE_LABEL.get(key, key)
        grade = "우수" if value >= 4 else "양호" if value >= 3 else "보통" if value >= 2 else "미흡"
        pdf.cell(60, 7, f"  {label}", border=1)
        pdf.cell(25, 7, f"{value}/5", border=1, align="C")
        pdf.cell(25, 7, grade, border=1, align="C")
        pdf.ln()
    pdf.ln(5)

    # 항목별 상세 평가
    if score_details:
        for detail in score_details:
            item_name = detail.get("item", "")
            item_score = detail.get("score", "")
            item_analysis = detail.get("analysis", "")
            pdf.set_font(fn, "B", 10)
            pdf.cell(0, 7, f"{item_name}: {item_score}", ln=True)
            pdf.set_font(fn, "", 9)
            pdf.multi_cell(0, 5, item_analysis)
            pdf.ln(3)
    separator()

    # ── 5. 핵심 강점 ──
    if top_strengths:
        heading("5. 핵심 강점 (Top 3)")
        for i, s in enumerate(top_strengths[:3], 1):
            pdf.set_font(fn, "B", 10)
            pdf.cell(0, 7, f"  {i}.", ln=False)
            pdf.set_font(fn, "", 9)
            pdf.ln()
            pdf.multi_cell(0, 5, f"    {s}")
            pdf.ln(2)
        separator()
    elif findings:
        heading("5. 핵심 발견")
        for i, f in enumerate(findings, 1):
            body(f"  {i}. {f}")
            pdf.ln(1)
        separator()

    # ── 6. 주요 우려사항 ──
    if top_concerns:
        heading("6. 주요 우려사항 (Top 3)")
        for i, c in enumerate(top_concerns[:3], 1):
            pdf.set_font(fn, "B", 10)
            pdf.cell(0, 7, f"  {i}.", ln=False)
            pdf.set_font(fn, "", 9)
            pdf.ln()
            pdf.multi_cell(0, 5, f"    {c}")
            pdf.ln(2)
        separator()

    # ── 7. 추가 확인 필요사항 ──
    if verification_needed:
        heading("7. 추가 확인 필요사항")
        for i, v in enumerate(verification_needed, 1):
            body(f"  {i}. {v}")
            pdf.ln(1)
        separator()

    # ── 8. 종합 의견 ──
    if overall_opinion:
        heading("8. 종합 의견")
        body(overall_opinion)
        separator()
    elif main_text:
        heading("8. AI 분석 의견")
        body(main_text)
        separator()

    # ── 부록: 외부 시장 조사 원본 ──
    if web_research:
        heading("부록: 외부 시장 조사 원본 (Tavily 웹 검색)")
        pdf.set_font(fn, "", 8)
        pdf.set_text_color(60, 60, 60)
        # 섹션별로 분리하여 출력
        for line in web_research.split("\n"):
            stripped = line.strip()
            if stripped.startswith("### "):
                pdf.ln(3)
                pdf.set_font(fn, "B", 9)
                pdf.cell(0, 6, stripped.replace("### ", ""), ln=True)
                pdf.set_font(fn, "", 8)
            elif stripped:
                pdf.multi_cell(0, 4, stripped)
        pdf.set_text_color(0, 0, 0)
        separator()

    # ── 규칙 기반 분석 (참고) ──
    if rule_part:
        heading("참고: 규칙 기반 분석")
        pdf.set_font(fn, "", 8)
        pdf.set_text_color(80, 80, 80)
        pdf.multi_cell(0, 5, rule_part)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)

    # ── 면책 조항 ──
    pdf.ln(5)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    pdf.set_font(fn, "", 7)
    pdf.set_text_color(150, 150, 150)
    pdf.multi_cell(0, 4, (
        "본 보고서는 AI에 의해 자동 생성된 참고 자료이며, 최종 투자 판단을 대체하지 않습니다. "
        "실제 투자 의사결정 시 추가 실사(Due Diligence)를 수행하시기 바랍니다. "
        f"보고서 ID: {analysis.id}"
    ))

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ── DOCX 생성 ─────────────────────────────────────────────────────

def generate_docx(analysis: AIAnalysis, startup: Startup) -> bytes:
    """AI 분석 결과를 풀 보고서 DOCX로 생성"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "맑은 고딕"

    type_label = TYPE_LABEL.get(analysis.analysis_type, analysis.analysis_type)
    rec_label = REC_LABEL.get(analysis.recommendation or "", analysis.recommendation or "-")
    risk_label = RISK_LABEL.get(analysis.risk_level or "", analysis.risk_level or "-")
    created = analysis.created_at.strftime("%Y.%m.%d") if analysis.created_at else datetime.now().strftime("%Y.%m.%d")

    scores = analysis.scores or {}
    main_text, findings, rule_part = _parse_sections(analysis.summary or "")

    full = _extract_full_report(analysis)
    score_details = full.get("score_details", []) if full else []
    top_strengths = full.get("top_strengths", []) if full else []
    top_concerns = full.get("top_concerns", []) if full else []
    verification_needed = full.get("verification_needed", []) if full else []
    overall_opinion = full.get("overall_opinion", "") if full else ""
    company_overview = full.get("company_overview", {}) if full else {}
    eligibility = full.get("eligibility_check", {}) if full else {}
    one_line = full.get("one_line_summary", "") if full else ""
    web_research = scores.get("_web_research", "")

    score_items = {k: v for k, v in scores.items() if isinstance(v, (int, float))}
    avg_score = sum(score_items.values()) / len(score_items) if score_items else 0
    sorted_scores = sorted(score_items.items(), key=lambda x: x[1], reverse=True)

    # ── 표지 ──
    title = doc.add_heading(f"{type_label} 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{startup.company_name}  |  {created}  |  {startup.industry or '-'}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(100, 100, 100)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = note.add_run("본 보고서는 AI 에이전트의 분석을 기반으로 작성되었으며, 담당자의 검토를 거쳐 확정되어야 합니다.")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(150, 150, 150)

    # ── 1. 요약 ──
    doc.add_heading("1. 요약", level=1)
    t = doc.add_table(rows=2, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["종합 점수", "판정", "리스크", "매력도"]):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs:
                rn.bold = True
                rn.font.size = Pt(9)
    for i, v in enumerate([f"{avg_score:.1f}/5", rec_label, risk_label, f"{analysis.investment_attractiveness or '-'}/5"]):
        c = t.rows[1].cells[i]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if one_line:
        doc.add_paragraph(f"한줄 요약: {one_line}")
    elif main_text:
        doc.add_paragraph(main_text[:500])

    # ── 2. 적격성 ──
    if eligibility:
        doc.add_heading("2. 기본 적격성 검사", level=1)
        labels = {"legal_issue": "법적 결격사유", "duplicate_application": "중복 지원", "antisocial_model": "반사회적 사업", "portfolio_conflict": "포트폴리오 충돌"}
        et = doc.add_table(rows=len(labels) + 1, cols=2)
        et.style = "Light Grid Accent 1"
        et.rows[0].cells[0].text = "항목"
        et.rows[0].cells[1].text = "결과"
        for idx, (key, label) in enumerate(labels.items(), 1):
            et.rows[idx].cells[0].text = label
            et.rows[idx].cells[1].text = eligibility.get(key, "미확인")

    # ── 3. 기업 개요 ──
    if company_overview:
        doc.add_heading("3. 기업 개요", level=1)
        labels = {"company_name": "기업명", "representative": "대표자", "business_field": "사업 분야", "main_product": "주요 제품", "team_size": "팀 규모", "investment_request": "투자 요청", "government_rd": "정부 R&D"}
        ct = doc.add_table(rows=len([v for v in company_overview.values() if v]) + 1, cols=2)
        ct.style = "Light Grid Accent 1"
        ct.rows[0].cells[0].text = "항목"
        ct.rows[0].cells[1].text = "내용"
        row_idx = 1
        for key, val in company_overview.items():
            if val and row_idx < len(ct.rows):
                ct.rows[row_idx].cells[0].text = labels.get(key, key)
                ct.rows[row_idx].cells[1].text = str(val)
                row_idx += 1

    # ── 4. 항목별 평가 ──
    doc.add_heading("4. 항목별 점수", level=1)
    st = doc.add_table(rows=len(sorted_scores) + 1, cols=3)
    st.style = "Light Grid Accent 1"
    for i, h in enumerate(["평가 항목", "점수", "등급"]):
        st.rows[0].cells[i].text = h
        for p in st.rows[0].cells[i].paragraphs:
            for rn in p.runs:
                rn.bold = True
    for idx, (key, value) in enumerate(sorted_scores, 1):
        grade = "우수" if value >= 4 else "양호" if value >= 3 else "보통" if value >= 2 else "미흡"
        st.rows[idx].cells[0].text = SCORE_LABEL.get(key, key)
        st.rows[idx].cells[1].text = f"{value}/5"
        st.rows[idx].cells[2].text = grade

    if score_details:
        doc.add_paragraph()
        for detail in score_details:
            p = doc.add_paragraph()
            r = p.add_run(f"{detail.get('item', '')}: {detail.get('score', '')}")
            r.bold = True
            r.font.size = Pt(10)
            doc.add_paragraph(detail.get("analysis", ""))

    # ── 5. 핵심 강점 ──
    if top_strengths:
        doc.add_heading("5. 핵심 강점 (Top 3)", level=1)
        for i, s in enumerate(top_strengths[:3], 1):
            doc.add_paragraph(f"{i}. {s}")
    elif findings:
        doc.add_heading("5. 핵심 발견", level=1)
        for i, f in enumerate(findings, 1):
            doc.add_paragraph(f"{i}. {f}")

    # ── 6. 주요 우려사항 ──
    if top_concerns:
        doc.add_heading("6. 주요 우려사항 (Top 3)", level=1)
        for i, c in enumerate(top_concerns[:3], 1):
            doc.add_paragraph(f"{i}. {c}")

    # ── 7. 추가 확인 필요사항 ──
    if verification_needed:
        doc.add_heading("7. 추가 확인 필요사항", level=1)
        for i, v in enumerate(verification_needed, 1):
            doc.add_paragraph(f"{i}. {v}")

    # ── 8. 종합 의견 ──
    if overall_opinion:
        doc.add_heading("8. 종합 의견", level=1)
        doc.add_paragraph(overall_opinion)
    elif main_text:
        doc.add_heading("8. AI 분석 의견", level=1)
        doc.add_paragraph(main_text)

    # ── 부록: 외부 시장 조사 원본 ──
    if web_research:
        doc.add_heading("부록: 외부 시장 조사 원본 (Tavily 웹 검색)", level=1)
        for line in web_research.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                p = doc.add_paragraph()
                rn = p.add_run(stripped.replace("### ", ""))
                rn.bold = True
                rn.font.size = Pt(10)
            else:
                p = doc.add_paragraph(stripped)
                for rn in p.runs:
                    rn.font.size = Pt(8)
                    rn.font.color.rgb = RGBColor(80, 80, 80)

    # ── 규칙 기반 ──
    if rule_part:
        doc.add_heading("참고: 규칙 기반 분석", level=1)
        p = doc.add_paragraph(rule_part)
        for rn in p.runs:
            rn.font.color.rgb = RGBColor(100, 100, 100)
            rn.font.size = Pt(9)

    # ── 면책 ──
    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        "본 보고서는 AI에 의해 자동 생성된 참고 자료이며, 최종 투자 판단을 대체하지 않습니다. "
        f"보고서 ID: {analysis.id}"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
